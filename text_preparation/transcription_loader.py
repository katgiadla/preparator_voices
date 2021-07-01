from itertools import chain
from docx import Document
from exceptions import EmptyListError
from json import dumps
from os.path import join, exists

def open_docx_transcription(path: str) -> list:
    """

    Opening docx file and save each one word of transcription to list.
    If document doesn't exist is thrown FileNotFoundException.

    Args:
        path (str): path to docx file

    Returns:
        transcription_list (list):  list of transcritptions' paragraphs

    """
    transcription_list = []

    try:
        transcript = Document(path)
    except:
        transcript = Document()

    for paragraph in transcript.paragraphs:
        transcription_list.append(paragraph.text)

    if len(transcription_list) == 0:
        raise FileNotFoundError("Document doesn't exist!")

    return transcription_list

def open_txt_transcription(path: str) -> list:
    """
    Opening txt file and save each one word of transcription to list.
    If document doesn't exist is thrown FileNotFoundException.

    Args:
        path (str): path to txt file

    Returns:
        transcription_list (list):  list of transcritptions' lines

    """
    transcription_list = []
    try:
        with open(path, 'r', encoding='utf-8') as file:
            transcription_list = file.readlines()
        file.close()
    except FileNotFoundError:
        print("File doesn't exist!")
        raise

    return transcription_list

def split_lines_to_words(transcriptions_lines: list) -> list:
    """
    Divide lines of transcriptions by the words
    If list of lines is empty, EmptyListError has been thrown
    Args:
        transcriptions_lines (list): lines of transcription

    Returns:
        words (list):   list of separated words
    """
    words = []

    if not transcriptions_lines:
        raise EmptyListError.EmptyListError
    else:
        for counter, line in enumerate(transcriptions_lines):
            words.append(line.split(' '))

    words = list(chain(*words))
    return words

def cleaning_transcription_list(transcription_words: list) -> list:
    """
    Cleaning from whitespaces and preparing list of words from transcriptions.
    If transcription is empty function throws EmptyListError

    Args:
        transcription_words (list):  list of words in documents

    Returns:
        transcript_words (list):    clean words from transcription
    """
    if not transcription_words:
        raise EmptyListError.EmptyListError
    else:
        for counter, word in enumerate(transcription_words):
            transcription_words[counter] = str(word).strip('\n\t\.\,')

    return transcription_words

def save_list_of_words_to_file(path_of_transcribed_sound: str, list_of_words: None, path_to_save: str) -> None:
    """
    Save transcribed list of words in JSON file.
    If list doesn't exist or list is empty, function throws exception.
    If file doesn't exist, function create a new.
    If directory doesn't exist, function throws exception.
    If directory doesn't exist, function throws exception.

    Args:
        path_of_transcribed_sound (str):    path to file, which be transcribed
        list_of_words (list):               list of transcribed words
        path_to_save (str):                 directory, where file should be saved

    Returns:
        None
    """
    if list_of_words == None or len(list_of_words) == 0:
        raise EmptyListError.EmptyListError("List of transcription doesn't exist or is empty")
    else:
        pass

    if exists(path_to_save) == False:
        raise NotADirectoryError("This director doesn't exist!")
    else:
        pass

    if exists(path_of_transcribed_sound) == False:
        raise FileNotFoundError("This file doesn't exist!")
    else:
        pass

    data_to_save = {'file_name': path_of_transcribed_sound,
                    'transcription': list_of_words}
    path = join(path_to_save, path_of_transcribed_sound + '.json')

    JSON_string = dumps(data_to_save)
    JSON_file = open(path, "w")
    JSON_file.write(JSON_string)
    JSON_file.close()