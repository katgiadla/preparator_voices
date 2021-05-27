from itertools import chain
from docx import Document
from exceptions import EmptyListError

def open_docx_transcription(path: str) -> list:
    """

    Opening docx file and save each one word of transcription to list.
    If document doesn't exist is thrown FileNotFoundException.

    Args:
        path (str): path to docx file

    Returns:
        transcription_list (list):  list of transcritptions' paragraphs

    """
    transcription_list = []

    try:
        transcript = Document(path)
    except:
        transcript = Document()

    for paragraph in transcript.paragraphs:
        transcription_list.append(paragraph.text)

    if len(transcription_list) == 0:
        raise FileNotFoundError("Document doesn't exist!")

    return transcription_list

def open_txt_transcription(path: str) -> list:
    """
    Opening txt file and save each one word of transcription to list.
    If document doesn't exist is thrown FileNotFoundException.

    Args:
        path (str): path to txt file

    Returns:
        transcription_list (list):  list of transcritptions' lines

    """
    transcription_list = []
    try:
        with open(path, 'r', encoding='utf-8') as file:
            transcription_list = file.readlines()
        file.close()
    except FileNotFoundError:
        print("File doesn't exist!")
        raise

    return transcription_list

def split_lines_to_words(transcriptions_lines: list) -> list:
    """
    Divide lines of transcriptions by the words
    If list of lines is empty, EmptyListError has been thrown
    Args:
        transcriptions_lines (list): lines of transcription

    Returns:
        words (list):   list of separated words
    """
    words = []

    if not transcriptions_lines:
        raise EmptyListError.EmptyListError
    else:
        for counter, line in enumerate(transcriptions_lines):
            words.append(line.split(' '))

    words = list(chain(*words))
    return words

def cleaning_transcription_list(transcription_words: list) -> list:
    """
    Cleaning from whitespaces and preparing list of words from transcriptions.
    If transcription is empty function throws EmptyListError

    Args:
        transcription_words (list):  list of words in documents

    Returns:
        transcript_words (list):    clean words from transcription
    """
    if not transcription_words:
        raise EmptyListError.EmptyListError
    else:
        for counter, word in enumerate(transcription_words):
            transcription_words[counter] = str(word).strip('\n\t\.\,')

    return transcription_words
